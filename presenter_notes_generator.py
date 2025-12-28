# presenter_notes_generator.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


def create_brick_breaker_presenter_notes():
    """Create detailed presenter notes for Brick Breaker game"""
    doc = Document()

    # Title page
    title = doc.add_heading('BRICK BREAKER GAME', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Video Demonstration - Presenter Notes', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    meta = doc.add_paragraph()
    meta.add_run('Game: Brick Breaker - Defender of the Crystal Kingdom\n')
    meta.add_run('Created: ' + datetime.datetime.now().strftime("%B %d, %Y") + '\n')
    meta.add_run('Duration: 3-5 minutes\n')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Introduction Section
    doc.add_heading('INTRODUCTION SCRIPT', 1)

    intro = doc.add_paragraph()
    intro.add_run("Opening Script:\n").bold = True
    intro.add_run(
        "\"Hello and welcome to our Brick Breaker game demonstration. "
        "Today, I'll be showcasing 'Defender of the Crystal Kingdom' - "
        "a modern take on the classic brick-breaking game with enhanced "
        "features, visual effects, and a compelling narrative.\"\n\n"
    )

    # Demonstration Outline
    doc.add_heading('DEMONSTRATION OUTLINE', 2)
    outline_items = [
        "1. Game Overview and Story (30 seconds)",
        "2. Main Menu and Controls (30 seconds)",
        "3. Core Gameplay Mechanics (60 seconds)",
        "4. Special Features and Power-ups (60 seconds)",
        "5. Level Progression (30 seconds)",
        "6. Technical Implementation (30 seconds)",
        "7. Q&A (30 seconds)"
    ]

    for item in outline_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Detailed Scene-by-Scene Notes
    doc.add_heading('SCENE-BY-SCENE PRESENTER NOTES', 1)

    scenes = [
        {
            'title': 'SCENE 1: GAME INTRODUCTION',
            'duration': '30 seconds',
            'actions': [
                'Open the game application',
                'Show game title screen',
                'Explain the narrative: "Defender of the Crystal Kingdom"',
                'Briefly mention the coursework context'
            ],
            'script': (
                "\"This is our Brick Breaker game developed for the "
                "Computer Game Development coursework. We've created a "
                "complete gaming experience with a story about defending "
                "crystal barriers from dark energy orbs.\""
            )
        },
        {
            'title': 'SCENE 2: MAIN MENU DEMONSTRATION',
            'duration': '30 seconds',
            'actions': [
                'Navigate through main menu options',
                'Show Play, Instructions, and Quit buttons',
                'Demonstrate button hover effects',
                'Enter Instructions screen'
            ],
            'script': (
                "\"Here's our main menu with clean, responsive design. "
                "Notice the hover effects and consistent color scheme. "
                "The instructions screen details all controls and gameplay mechanics.\""
            )
        },
        {
            'title': 'SCENE 3: GAME CONTROLS',
            'duration': '30 seconds',
            'actions': [
                'Show control options (mouse and keyboard)',
                'Demonstrate paddle movement',
                'Launch ball with spacebar',
                'Show pause functionality'
            ],
            'script': (
                "\"Players can use either mouse movement or arrow keys "
                "for paddle control. Pressing spacebar launches the ball, "
                "and P pauses the game. This dual-control system makes "
                "the game accessible to different players.\""
            )
        },
        {
            'title': 'SCENE 4: CORE GAMEPLAY - LEVEL 1',
            'duration': '60 seconds',
            'actions': [
                'Start Level 1',
                'Demonstrate basic brick breaking',
                'Show collision physics',
                'Explain scoring system',
                'Break several bricks to show particle effects'
            ],
            'script': (
                "\"Starting with Level 1, we have simple brick patterns. "
                "Watch how the ball physics work - angles change based on "
                "where the ball hits the paddle. Each brick break gives "
                "points and creates particle effects for visual feedback.\""
            )
        },
        {
            'title': 'SCENE 5: BRICK TYPES AND POWER-UPS',
            'duration': '60 seconds',
            'actions': [
                'Break different brick types (normal, tough, power-up)',
                'Collect a power-up',
                'Demonstrate multi-ball effect',
                'Show paddle extension power-up',
                'Explain combo scoring system'
            ],
            'script': (
                "\"We have four brick types: normal, tough, power-up, "
                "and unbreakable. Power-up bricks drop special abilities "
                "like multi-ball or paddle extension. Notice the combo "
                "system - breaking bricks quickly gives bonus points.\""
            )
        },
        {
            'title': 'SCENE 6: LEVEL PROGRESSION',
            'duration': '45 seconds',
            'actions': [
                'Complete Level 1',
                'Show level completion screen',
                'Start Level 2',
                'Demonstrate increased difficulty',
                'Show moving bricks in Level 3'
            ],
            'script': (
                "\"After breaking all bricks, we complete the level and "
                "get a bonus. Level 2 introduces tougher patterns, and "
                "Level 3 has moving bricks for added challenge. The game "
                "gets progressively harder to maintain engagement.\""
            )
        },
        {
            'title': 'SCENE 7: GAME OVER AND HIGH SCORES',
            'duration': '30 seconds',
            'actions': [
                'Intentionally lose all lives',
                'Show game over screen',
                'Demonstrate restart functionality',
                'Show high score tracking'
            ],
            'script': (
                "\"If you lose all lives, the game ends with your final "
                "score. High scores are saved locally. Players can restart "
                "or return to the main menu to try again.\""
            )
        },
        {
            'title': 'SCENE 8: TECHNICAL FEATURES',
            'duration': '45 seconds',
            'actions': [
                'Switch to code view (if recording code)',
                'Show modular architecture',
                'Explain collision detection algorithm',
                'Show particle system implementation',
                'Mention Python and Pygame usage'
            ],
            'script': (
                "\"Technically, we implemented this using Python and Pygame "
                "with a modular architecture. Each game component is a "
                "separate class for maintainability. The collision detection "
                "uses precise algorithms, and our particle system adds "
                "polish to the gameplay experience.\""
            )
        },
        {
            'title': 'SCENE 9: CONCLUSION',
            'duration': '30 seconds',
            'actions': [
                'Return to game view',
                'Show all features working together',
                'End with game running smoothly'
            ],
            'script': (
                "\"In conclusion, our Brick Breaker game demonstrates "
                "complete game development from concept to implementation. "
                "It includes all required features plus advanced elements "
                "like particle effects, power-ups, and a progression system. "
                "Thank you for watching!\""
            )
        }
    ]

    for i, scene in enumerate(scenes):
        # Add scene header
        scene_header = doc.add_heading(f'{scene["title"]} - {scene["duration"]}', 2)
        scene_header.paragraph_format.space_before = Pt(12)

        # Add actions
        actions_para = doc.add_paragraph()
        actions_para.add_run('Actions to Demonstrate:\n').bold = True

        for action in scene['actions']:
            doc.add_paragraph(f'• {action}', style='List Bullet')

        # Add script
        script_para = doc.add_paragraph()
        script_para.add_run('Presenter Script:\n').bold = True
        script_para.add_run(scene['script'])

        # Add spacing between scenes
        doc.add_paragraph()

        # Add page break after every 2 scenes
        if (i + 1) % 2 == 0 and i < len(scenes) - 1:
            doc.add_page_break()

    # Add Technical Details Appendix
    doc.add_page_break()
    doc.add_heading('TECHNICAL IMPLEMENTATION DETAILS', 1)

    tech_details = [
        "Game Engine: Pygame 2.5.1",
        "Programming Language: Python 3.8+",
        "Architecture: Object-Oriented with modular design",
        "File Structure:",
        "  - main.py: Game entry point",
        "  - game.py: Main game controller",
        "  - paddle.py: Paddle class with movement logic",
        "  - ball.py: Ball physics and collision",
        "  - bricks.py: Brick management system",
        "  - powerups.py: Power-up effects",
        "  - particles.py: Visual effects system",
        "  - ui.py: User interface components",
        "  - config.py: Game constants and settings",
        "",
        "Key Algorithms:",
        "  - Collision Detection: Advanced side-detection algorithm",
        "  - Physics: Realistic ball bouncing with angle calculation",
        "  - Particle System: Efficient particle management",
        "  - State Management: Game state machine for menus/gameplay",
        "",
        "Performance:",
        "  - Target FPS: 60 frames per second",
        "  - Resolution: 800x600 pixels",
        "  - Memory Management: Efficient object pooling",
        "",
        "Features Implemented:",
        "  ✓ Modular game architecture",
        "  ✓ Multiple brick types with different behaviors",
        "  ✓ Power-up system with 4 different abilities",
        "  ✓ Combo scoring system",
        "  ✓ Particle effects for visual feedback",
        "  ✓ 3 levels with progressive difficulty",
        "  ✓ High score tracking",
        "  ✓ Responsive UI with hover effects",
        "  ✓ Pause functionality",
        "  ✓ Game state management"
    ]

    for detail in tech_details:
        if detail.endswith(':'):
            para = doc.add_paragraph()
            para.add_run(detail).bold = True
        elif detail.startswith('  - ') or detail.startswith('  ✓ '):
            doc.add_paragraph(detail)
        else:
            doc.add_paragraph(detail)

    # Save document
    filename = 'brick_breaker_presenter_notes.docx'
    doc.save(filename)
    print(f"Brick Breaker presenter notes saved to {filename}")

    return filename


def create_snake_game_presenter_notes():
    """Create detailed presenter notes for Snake game"""
    doc = Document()

    # Title page
    title = doc.add_heading('EMERALD SERPENT GAME', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Video Demonstration - Presenter Notes', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    meta = doc.add_paragraph()
    meta.add_run('Game: Emerald Serpent - The Mystical Garden Quest\n')
    meta.add_run('Created: ' + datetime.datetime.now().strftime("%B %d, %Y") + '\n')
    meta.add_run('Duration: 3-5 minutes\n')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Introduction Section
    doc.add_heading('INTRODUCTION SCRIPT', 1)

    intro = doc.add_paragraph()
    intro.add_run("Opening Script:\n").bold = True
    intro.add_run(
        "\"Welcome to our Snake game demonstration: 'Emerald Serpent - "
        "The Mystical Garden Quest.' This is a modern implementation of "
        "the classic Snake game with enhanced graphics, multiple levels, "
        "and special features that we developed for our coursework.\"\n\n"
    )

    # Demonstration Outline
    doc.add_heading('DEMONSTRATION OUTLINE', 2)
    outline_items = [
        "1. Game Introduction and Narrative (30 seconds)",
        "2. Main Menu and Instructions (30 seconds)",
        "3. Basic Gameplay Mechanics (60 seconds)",
        "4. Food System and Special Items (60 seconds)",
        "5. Obstacles and Level Design (45 seconds)",
        "6. Progression and Scoring (45 seconds)",
        "7. Technical Features (30 seconds)"
    ]

    for item in outline_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Detailed Scene-by-Scene Notes
    doc.add_heading('SCENE-BY-SCENE PRESENTER NOTES', 1)

    scenes = [
        {
            'title': 'SCENE 1: GAME INTRODUCTION',
            'duration': '30 seconds',
            'actions': [
                'Launch the game',
                'Show title screen with "Emerald Serpent"',
                'Explain the mystical garden narrative',
                'Mention coursework context'
            ],
            'script': (
                "\"This Snake game was developed for our Computer Game "
                "Development module. We've created a mystical garden theme "
                "where an emerald serpent seeks enchanted fruits. "
                "It's not just a simple Snake game - it has levels, "
                "obstacles, and special features.\""
            )
        },
        {
            'title': 'SCENE 2: MAIN MENU AND CONTROLS',
            'duration': '30 seconds',
            'actions': [
                'Navigate main menu options',
                'Enter instructions screen',
                'Show control scheme (arrow keys and WASD)',
                'Explain game objectives'
            ],
            'script': (
                "\"Our main menu provides clear options. The instructions "
                "detail all controls and gameplay rules. Players can use "
                "either arrow keys or WASD for movement. The objective is "
                "simple: eat food, grow longer, and avoid collisions.\""
            )
        },
        {
            'title': 'SCENE 3: STARTING GAMEPLAY - LEVEL 1',
            'duration': '45 seconds',
            'actions': [
                'Start Level 1',
                'Show snake initial position (3 segments)',
                'Demonstrate movement controls',
                'Eat first food item',
                'Show growth animation'
            ],
            'script': (
                "\"Starting Level 1, our snake begins with 3 segments. "
                "Notice the smooth grid-based movement and the visual "
                "gradient from head to tail. When we eat food, the snake "
                "grows and we earn points. Watch the particle effects "
                "when collecting food.\""
            )
        },
        {
            'title': 'SCENE 4: FOOD SYSTEM',
            'duration': '60 seconds',
            'actions': [
                'Show different food types (normal, golden, speed)',
                'Demonstrate golden food giving 2 growth segments',
                'Show speed food effect',
                'Explain food expiration timer',
                'Demonstrate combo system'
            ],
            'script': (
                "\"We have three food types: normal red food, golden "
                "food worth more points, and blue speed food. Golden food "
                "makes the snake grow by 2 segments. Speed food temporarily "
                "increases movement speed. Eating quickly builds a combo "
                "for score multiplier.\""
            )
        },
        {
            'title': 'SCENE 5: OBSTACLES AND LEVEL DESIGN',
            'duration': '45 seconds',
            'actions': [
                'Move to Level 2',
                'Show static obstacle patterns',
                'Navigate around obstacles',
                'Move to Level 3',
                'Show maze-like obstacle layout'
            ],
            'script': (
                "\"Level 2 introduces static obstacles - garden statues "
                "that block the snake's path. Level 3 has a maze-like "
                "pattern requiring careful navigation. Each level has "
                "different obstacle layouts to increase challenge.\""
            )
        },
        {
            'title': 'SCENE 6: SNAKE MECHANICS AND CHALLENGES',
            'duration': '45 seconds',
            'actions': [
                'Show increasing speed as snake grows',
                'Demonstrate self-collision avoidance',
                'Show wall collision mechanics',
                'Demonstrate tight space navigation'
            ],
            'script': (
                "\"As the snake grows longer, it also moves faster, "
                "increasing difficulty. The main challenge is avoiding "
                "self-collision - hitting your own body ends the game. "
                "Walls are also solid boundaries that can't be passed.\""
            )
        },
        {
            'title': 'SCENE 7: LEVEL PROGRESSION AND SCORING',
            'duration': '45 seconds',
            'actions': [
                'Reach target length for level completion',
                'Show level complete screen with bonus',
                'Demonstrate high score tracking',
                'Show game over screen with statistics'
            ],
            'script': (
                "\"To complete a level, reach the target length shown "
                "in the HUD. Success gives a level completion bonus. "
                "High scores are tracked locally. The game over screen "
                "shows detailed statistics of your performance.\""
            )
        },
        {
            'title': 'SCENE 8: VISUAL AND TECHNICAL FEATURES',
            'duration': '45 seconds',
            'actions': [
                'Show particle effects system',
                'Demonstrate snake animation (wiggle effect)',
                'Show gradient coloring on snake',
                'Explain modular code architecture',
                'Mention Python/Pygame implementation'
            ],
            'script': (
                "\"Technically, we implemented a particle system for "
                "visual effects, animated snake movement with a wiggle "
                "effect, and gradient coloring. The code uses a modular "
                "OOP architecture in Python with Pygame, making it "
                "maintainable and extensible.\""
            )
        },
        {
            'title': 'SCENE 9: CONCLUSION AND FEATURE SUMMARY',
            'duration': '30 seconds',
            'actions': [
                'Show all features working together',
                'Highlight key achievements',
                'End with smooth gameplay demonstration'
            ],
            'script': (
                "\"In summary, our Emerald Serpent game demonstrates "
                "modern game development with classic gameplay. We've "
                "implemented multiple levels, different food types, "
                "obstacles, visual effects, and a progression system - "
                "all while maintaining clean, modular code. Thank you!\""
            )
        }
    ]

    for i, scene in enumerate(scenes):
        # Add scene header
        scene_header = doc.add_heading(f'{scene["title"]} - {scene["duration"]}', 2)
        scene_header.paragraph_format.space_before = Pt(12)

        # Add actions
        actions_para = doc.add_paragraph()
        actions_para.add_run('Actions to Demonstrate:\n').bold = True

        for action in scene['actions']:
            doc.add_paragraph(f'• {action}', style='List Bullet')

        # Add script
        script_para = doc.add_paragraph()
        script_para.add_run('Presenter Script:\n').bold = True
        script_para.add_run(scene['script'])

        # Add spacing between scenes
        doc.add_paragraph()

        # Add page break after every 2 scenes
        if (i + 1) % 2 == 0 and i < len(scenes) - 1:
            doc.add_page_break()

    # Add Technical Details Appendix
    doc.add_page_break()
    doc.add_heading('TECHNICAL IMPLEMENTATION DETAILS', 1)

    tech_details = [
        "Game Engine: Pygame 2.5.1",
        "Programming Language: Python 3.8+",
        "Architecture: Object-Oriented with component-based design",
        "File Structure:",
        "  - main.py: Game entry point",
        "  - game.py: Main game controller",
        "  - snake.py: Snake class with movement and growth",
        "  - food.py: Food generation and types",
        "  - grid.py: Grid and obstacle management",
        "  - particles.py: Visual effects system",
        "  - ui.py: User interface components",
        "  - config.py: Game constants and settings",
        "",
        "Key Algorithms:",
        "  - Grid-based Movement: Precise tile-based navigation",
        "  - Collision Detection: Efficient position checking",
        "  - Food Spawning: Smart placement avoiding snake/obstacles",
        "  - Path Finding: (Optional) For advanced features",
        "",
        "Visual Features:",
        "  - Particle Effects: Food collection, collisions, trails",
        "  - Snake Animation: Gradient colors and wiggle movement",
        "  - Food Animation: Pulsing, sparkling effects",
        "  - Obstacle Design: 3D-like textures with cracks",
        "",
        "Game Features:",
        "  ✓ 3 levels with different obstacle patterns",
        "  ✓ 3 food types with different effects",
        "  ✓ Combo scoring system (1.5x multiplier)",
        "  ✓ Progressive speed increase",
        "  ✓ High score tracking",
        "  ✓ Level progression system",
        "  ✓ Visual feedback for all actions",
        "  ✓ Responsive controls",
        "  ✓ Game state management"
    ]

    for detail in tech_details:
        if detail.endswith(':'):
            para = doc.add_paragraph()
            para.add_run(detail).bold = True
        elif detail.startswith('  - ') or detail.startswith('  ✓ '):
            doc.add_paragraph(detail)
        else:
            doc.add_paragraph(detail)

    # Save document
    filename = 'snake_game_presenter_notes.docx'
    doc.save(filename)
    print(f"Snake game presenter notes saved to {filename}")

    return filename


def create_recording_guide():
    """Create a guide for recording gameplay videos"""
    doc = Document()

    # Title
    title = doc.add_heading('GAMEPLAY VIDEO RECORDING GUIDE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('For Coursework Submission', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Introduction
    doc.add_heading('1. PRE-RECORDING PREPARATION', 1)

    prep_steps = [
        "Test your game thoroughly to ensure no bugs during recording",
        "Close unnecessary applications to improve performance",
        "Clean up your desktop background",
        "Set game window to appropriate size (800x600 recommended)",
        "Practice your presentation using the presenter notes",
        "Check microphone and audio quality",
        "Ensure good lighting if using webcam overlay"
    ]

    for step in prep_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('2. RECORDING SETUP', 1)

    setup_steps = [
        "Position game window clearly on screen",
        "Open presenter notes window on secondary monitor or side-by-side",
        "Start the recording script: python video_recorder_with_notes.py --game brick_breaker",
        "Verify recording is working (red 'RECORDING' indicator will appear)",
        "Do a 5-second test recording and play it back"
    ]

    for step in setup_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('3. RECORDING TECHNIQUES', 1)

    techniques = [
        "Speak clearly and at a moderate pace",
        "Demonstrate features while explaining them",
        "Use the notes as guide, but speak naturally",
        "Show both successful gameplay and error cases",
        "Highlight technical features mentioned in coursework",
        "Keep each demonstration segment under 60 seconds",
        "Maintain consistent audio levels throughout"
    ]

    for technique in techniques:
        doc.add_paragraph(technique, style='List Bullet')

    doc.add_heading('4. CONTENT STRUCTURE', 1)

    content_table = doc.add_table(rows=1, cols=3)
    hdr_cells = content_table.rows[0].cells
    hdr_cells[0].text = 'Time'
    hdr_cells[1].text = 'Content'
    hdr_cells[2].text = 'Key Points'

    content_data = [
        ['0:00-0:30', 'Introduction', 'Game name, purpose, coursework context'],
        ['0:30-1:00', 'Main Menu', 'Show all options, instructions screen'],
        ['1:00-2:00', 'Basic Gameplay', 'Core mechanics, controls, HUD'],
        ['2:00-3:00', 'Advanced Features', 'Special items, power-ups, combos'],
        ['3:00-4:00', 'Progression', 'Levels, difficulty increase, scoring'],
        ['4:00-4:30', 'Technical Details', 'Code structure, algorithms used'],
        ['4:30-5:00', 'Conclusion', 'Feature summary, closing remarks']
    ]

    for time, content, points in content_data:
        row_cells = content_table.add_row().cells
        row_cells[0].text = time
        row_cells[1].text = content
        row_cells[2].text = points

    doc.add_heading('5. POST-RECORDING', 1)

    post_steps = [
        "Review the recording for technical issues",
        "Edit if necessary (trim beginning/end, remove mistakes)",
        "Add title screen if desired",
        "Export to appropriate format (MP4 recommended)",
        "Upload to YouTube as unlisted or public",
        "Test the YouTube link before submission",
        "Include link in coursework document"
    ]

    for step in post_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Save document
    filename = 'recording_guide.docx'
    doc.save(filename)
    print(f"Recording guide saved to {filename}")

    return filename


def main():
    """Main function to generate all documentation"""
    print("=" * 60)
    print("PRESENTER NOTES AND RECORDING TOOLS GENERATOR")
    print("=" * 60)

    print("\nGenerating documentation for both games...")

    # Create presenter notes
    print("\n1. Creating Brick Breaker presenter notes...")
    bb_notes = create_brick_breaker_presenter_notes()

    print("\n2. Creating Snake Game presenter notes...")
    snake_notes = create_snake_game_presenter_notes()

    print("\n3. Creating recording guide...")
    guide = create_recording_guide()

    print("\n" + "=" * 60)
    print("DOCUMENTATION GENERATED SUCCESSFULLY!")
    print("=" * 60)
    print(f"\nGenerated files:")
    print(f"1. {bb_notes} - Brick Breaker presenter notes")
    print(f"2. {snake_notes} - Snake Game presenter notes")
    print(f"3. {guide} - Recording guide")

    print("\n" + "=" * 60)
    print("HOW TO USE THESE FILES:")
    print("=" * 60)
    print("\nFor recording Brick Breaker:")
    print("1. python video_recorder_with_notes.py --game brick_breaker")
    print("2. Follow notes in brick_breaker_presenter_notes.docx")
    print("3. Use notes window to guide your presentation")

    print("\nFor recording Snake Game:")
    print("1. python video_recorder_with_notes.py --game snake")
    print("2. Follow notes in snake_game_presenter_notes.docx")
    print("3. Use notes window to guide your presentation")

    print("\nTips for successful recording:")
    print("- Practice with the notes first")
    print("- Speak clearly and demonstrate features")
    print("- Keep the video under 5 minutes")
    print("- Show both gameplay and technical aspects")


if __name__ == "__main__":
    # Install required packages if not already installed
    required_packages = ['pygame', 'opencv-python', 'pillow', 'pyautogui', 'python-docx']

    print("Checking required packages...")
    for package in required_packages:
        try:
            __import__(package.replace('-', '_'))
            print(f"✓ {package}")
        except ImportError:
            print(f"✗ {package} not installed")
            print(f"  Install with: pip install {package}")

    print("\nTo install all packages:")
    print("pip install pygame opencv-python pillow pyautogui python-docx")

    # Ask user if they want to generate documents
    response = input("\nGenerate presenter notes and guides? (y/n): ")
    if response.lower() == 'y':
        main()